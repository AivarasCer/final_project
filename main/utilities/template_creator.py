import os
from docx import Document
from docxcompose.composer import Composer

# context = {'translation_dir': '/Translation from Lithuanian/'}


def save_document(template_path, context, doc_name, output_dir):
    """Render and save DOCX template."""
    template = Document(template_path)
    template.render(context)
    output_filename = f'{doc_name}.docx'
    output_path = os.path.join(output_dir, output_filename)
    template.save(output_path)
    return output_path


def template_composer(doc_template_path, ocr_output_path, output_dir):
    """Combine template with OCR output"""
    main_template = Document(doc_template_path)
    composer = Composer(main_template)

    ocr_output = Document(ocr_output_path)
    composer.append(ocr_output)

    combined_path = os.path.join(output_dir, 'processed.docx')
    composer.save(combined_path)
    return combined_path

